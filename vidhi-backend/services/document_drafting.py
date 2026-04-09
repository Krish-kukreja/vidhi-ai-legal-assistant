import logging
import os
import uuid
import time
from typing import Dict, Any, Optional

import docx
from docx.shared import Pt, Inches

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_aws import ChatBedrock

from configs import config
from stores.chroma import load_vectorstore, create_retriever

logger = logging.getLogger(__name__)

CROSS_QUESTION_PROMPT = ChatPromptTemplate.from_template("""\
You are an expert yet accessible AI legal assistant for VIDHI. A user from the general public wants to create a {document_type}.
So far, they have provided the following information:
Parties: {parties}
Key Terms: {key_terms}

Identify 1-3 critical missing details required to draft a legally sound standard {document_type} in India.
Ask the user a clear, plain-English question to gather this missing information. Do not use legal jargon. 
If no critical information is missing, output 'READY'.
Question:
""")

DRAFTING_PROMPT = ChatPromptTemplate.from_template("""\
You are an expert Indian corporate and family lawyer acting as an AI legal draftsman for VIDHI. 
Your task is to draft a legal document based on user requirements.

You have been provided a TEMPLATE for this document.
{persona_instructions}

TEMPLATE (Reference this heavily):
{template_context}

USER REQUIREMENTS:
Document Type: {document_type}
Parties Involved: {parties}
Key Terms & Specific Conditions: {key_terms}

INSTRUCTIONS:
1. Provide ONLY the finalized draft text. 
2. DO NOT include conversational filler like "Here is your drafted agreement..."
3. Use Markdown formatting (headers with #, subheaders with ##, bolding for party names).
4. Do not leave placeholder brackets [ ] if the user provided the information.
5. If the user missed critical standard information (e.g., date), use blank lines _____ or [Date] for them to fill later.
6. The jurisdiction is India unless specified otherwise.

FINAL LEGAL DRAFT:
""")

class DocumentDraftingService:
    def __init__(self):
        self._logger = logger
        self.embeddings = config.get_embeddings()
        
        # We load a dedicated vector store for templates
        if self.embeddings:
            self.vectorstore = load_vectorstore(self.embeddings, collection_name="vidhi-templates")
            self.retriever = create_retriever(self.vectorstore) if self.vectorstore else None
        else:
            self.vectorstore = None
            self.retriever = None

        # Standard LLM (Sonnet or Haiku)
        self.llm = ChatBedrock(
            model_id=config.BEDROCK_MODEL_ID_ADVANCED,
            region_name=config.AWS_REGION,
            model_kwargs={
                "temperature": 0.2, # Keep hallucination low for legal drafts
                "max_tokens": 4096
            }
        )
        
        self.cross_question_chain = CROSS_QUESTION_PROMPT | self.llm | StrOutputParser()
        self.drafting_chain = DRAFTING_PROMPT | self.llm | StrOutputParser()
        
        # Ensure output directory exists for docx files
        self.output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "storage", "drafts")
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_cross_questions(self, document_type: str, parties: str, key_terms: str) -> Dict[str, Any]:
        """
        AI Cross-questioning: Determines if more info is needed from a general user
        before drafting the document.
        """
        try:
            self._logger.info("Generating cross questions for general user...")
            question = self.cross_question_chain.invoke({
                "document_type": document_type,
                "parties": parties,
                "key_terms": key_terms
            })
            
            is_ready = "READY" in question.upper()
            return {
                "success": True,
                "needs_more_info": not is_ready,
                "question": None if is_ready else question.strip()
            }
        except Exception as e:
            self._logger.error(f"Error in cross questioning: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def generate_draft(self, document_type: str, parties: str, key_terms: str, 
                       persona: str = 'general', custom_template: Optional[str] = None) -> Dict[str, Any]:
        """
        Retrieves standard or custom template and generates a filled out draft using LLM.
        Supports personas: 'general' (standard fallback) and 'professional' (strict adherence to custom template).
        """
        try:
            self._logger.info(f"Generating draft for type: {document_type} (Persona: {persona})")
            
            # Retrieve relevant templates
            if custom_template:
                template_context = custom_template
                self._logger.info("Using provided custom firm template.")
            else:
                template_context = "No standard template found. Rely on general legal knowledge."
                if self.retriever:
                    docs = self.retriever.invoke(f"{document_type} template")
                    if docs:
                        # Just take the best matching template
                        template_context = docs[0].page_content
                        self._logger.info(f"Retrieved standard template context: {template_context[:50]}...")
            
            # Persona Instructions
            if persona == 'professional':
                persona_instructions = (
                    "PERSONA: Professional Lawyer.\n"
                    "STRICT REQUIREMENT: You MUST strictly adhere to the exact structure, proprietary language, "
                    "and boilerplate clauses ('and all' specific nuances) found in the TEMPLATE. Only extrapolate "
                    "and insert the specific dynamic variables provided in the USER REQUIREMENTS."
                )
            else:
                persona_instructions = (
                    "PERSONA: General Public User.\n"
                    "Use the TEMPLATE as a standard structural foundation, but translate complex requirements "
                    "provided by the user into logically sound standard clauses. Ensure the final document "
                    "is comprehensive and robust."
                )
            
            # Generate the draft
            markdown_draft = self.drafting_chain.invoke({
                "persona_instructions": persona_instructions,
                "template_context": template_context,
                "document_type": document_type,
                "parties": parties,
                "key_terms": key_terms
            })
            
            # Create a docx file
            docx_filename = f"{document_type.replace(' ', '_')}_{int(time.time())}.docx"
            docx_path = os.path.join(self.output_dir, docx_filename)
            
            self._convert_markdown_to_docx(markdown_draft, docx_path)
            
            return {
                "success": True,
                "markdown_draft": markdown_draft,
                "download_url": f"/api/v1/documents/download/{docx_filename}",
                "template_used": "Custom" if custom_template else ("Standard" if template_context != "No standard template found. Rely on general legal knowledge." else "None")
            }
            
        except Exception as e:
            self._logger.error(f"Error generating legal draft: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }

    def _convert_markdown_to_docx(self, markdown_text: str, file_path: str):
        """
        A simple converter to turn markdown into a basic docx file.
        In a production environment, you might use a more robust md->docx converter.
        """
        doc = docx.Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        lines = markdown_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('# '):
                p = doc.add_paragraph(line[2:])
                p.style = doc.styles['Heading 1']
            elif line.startswith('## '):
                p = doc.add_paragraph(line[3:])
                p.style = doc.styles['Heading 2']
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                p.style = doc.styles['Heading 3']
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
        
        # Save document
        doc.save(file_path)

# Singleton instance
document_drafting_service = DocumentDraftingService()
